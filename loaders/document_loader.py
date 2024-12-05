
      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      

      # docparseai/loaders/document_loader.py

import os
from typing import Optional
from pathlib import Path

from PyPDF2 import PdfReader
import docx


class DocumentLoader:
    @staticmethod
    def load(file_path: str) -> str:
        ext = Path(file_path).suffix.lower()

        if ext == '.pdf':
            return DocumentLoader._load_pdf(file_path)
        elif ext == '.docx':
            return DocumentLoader._load_docx(file_path)
        elif ext == '.txt':
            return DocumentLoader._load_txt(file_path)
        else:
            raise ValueError(fUnsupported file format: {ext})

    @staticmethod
    def _load_pdf(file_path: str) -> str:
        reader = PdfReader(file_path)
        text = []
        for page in reader.pages:
            text.append(page.extract_text() or )
        return n.join(text)

    @staticmethod
    def _load_docx(file_path: str) -> str:
        doc = docx.Document(file_path)
        return n.join([para.text for para in doc.paragraphs])

    @staticmethod
    def _load_txt(file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()

      
      
